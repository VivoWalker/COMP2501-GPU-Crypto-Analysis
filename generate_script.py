"""
Generate presentation script .docx from script prompt.txt specifications.
Output: presentation/presentation_script.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent
OUT_PATH = PROJECT_ROOT / "presentation" / "presentation_script.docx"

doc = Document()

# --- Style setup ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# --- Script content: (slide_label, speaker_notes) ---
# Target: ~1350 words, ~9.5 minutes at 140 wpm
# Slides 8-16 get proportionally more words

slides = [
    # ===== SECTION 1: INTRODUCTION (Slides 1-7, ~2 min, ~300 words) =====
    (
        "Slide 1 — Title Slide",
        "Good morning everyone. Today I'll present our analysis of NVIDIA GPU market "
        "dynamics and cryptocurrency correlations — a full-cycle study from the mining "
        "boom of 2020 through to the crash of 2022."
    ),

    (
        "Slide 2 — Background",
        "Here's what happened. From late 2020, Ethereum surged from about four hundred "
        "dollars to nearly four thousand eight hundred. This made GPU mining extremely "
        "profitable. Miners bought cards in bulk, causing global shortages, and eBay "
        "prices hit two to three times MSRP — the official retail price.\n\n"
        "In May 2021, NVIDIA fought back with LHR — Lite Hash Rate technology — that "
        "halved mining performance on new cards. Then in mid-2022, crypto crashed. "
        "Ethereum fell to around thirteen hundred dollars, and GPU secondary prices "
        "collapsed with it. Our core question: was this relationship real and measurable?"
    ),

    (
        "Slide 3 — Research Questions",
        "We asked four questions. One: do crypto price movements actually drive GPU "
        "secondary-market prices? Two: which GPU attributes — VRAM, hash rate, LHR — "
        "explain the size of price premiums over MSRP? Three: do different GPU tiers "
        "show different price volatility? And four: did LHR actually reduce miners' "
        "willingness to pay, and how persistent was the effect?"
    ),

    (
        "Slide 4 — Research Challenges",
        "We faced two main challenges. First, multicollinearity — cards with more VRAM "
        "also have higher MSRP and higher hash rates. These variables are tangled "
        "together, making it hard to isolate individual effects. Second, a frequency "
        "mismatch — crypto data is daily, but GPU eBay data is monthly. We had to "
        "align them carefully."
    ),

    (
        "Slide 5 — Data Sources",
        "We used three data sources. Cryptocurrency prices from the CoinGecko API — "
        "daily BTC and ETH across hundreds of exchanges. GPU specs — VRAM, hash rate, "
        "MSRP — from NVIDIA's official documentation.\n\n"
        "For GPU eBay prices, there's no public API — so we used AI-driven text mining "
        "to extract data from multiple published sources. Specifically, three tech "
        "journalism outlets: Tom's Hardware, PCMag, and 3D Center. All three "
        "independently tracked eBay sold listings, filtering out box-only listings, "
        "broken cards, and bundles — so we're looking at genuine sold prices only."
    ),

    (
        "Slide 6 — Tools and Software",
        "We used Python with pandas, statsmodels, and scipy for data processing and "
        "regression modeling; matplotlib and seaborn for visualization. R with ggplot2 "
        "was used for several figures. Claude Code assisted with web scraping scripts "
        "for the more complex data collection tasks."
    ),

    (
        "Slide 7 — Final Dataset",
        "Our final dataset: ten GPU models — seven non-LHR and three LHR variants — "
        "across twenty-four months, giving us over one hundred ninety data points. "
        "Critically, every single data point is from a real published source. Zero "
        "fabrication."
    ),

    # ===== SECTION 2: RESULTS — Slides 8-16 (~5.5 min, ~770 words) =====

    (
        "Slide 8 — Figure 1: Time Series Overview",
        "Let's start with the big picture. This chart overlays ETH price with GPU "
        "eBay prices for each tier — XX60 through XX90.\n\n"
        "The key observation: GPU prices and ETH are visually highly synchronized. "
        "They rise together, peak around May 2021 when LHR launched, and crash "
        "together through mid-2022. This holds across ALL tiers.\n\n"
        "The bottom panel shows individual models. The XX90 hit nearly thirty-two "
        "hundred dollars at peak; the XX60 peaked around nine hundred. The sync "
        "is clear — but we need to measure how strong the link actually is."
    ),

    (
        "Slide 9 — Figure 2: How Strong Is the Link?",
        "To quantify the link, we use two panels. Left: ETH price versus GPU price. "
        "Clear positive correlation, but there's a problem — both share a common "
        "boom-bust trend, which can inflate the correlation.\n\n"
        "So on the right, we use first-differencing — we compare monthly changes, "
        "not price levels. This removes the shared trend. Stars mark statistically "
        "significant correlations.\n\n"
        "The result: it IS quite strong. RTX 3090 shows r equals zero point five "
        "six, p below zero point zero one. Crypto prices do drive GPU prices — "
        "this is real, not a statistical artifact."
    ),

    (
        "Slide 10 — Premium Drivers: VRAM?",
        "Now — what drives the price premium, the percentage markup over MSRP?\n\n"
        "Our first guess: VRAM. More memory should mean better mining, so higher "
        "premiums. This chart plots premium against VRAM with the red regression "
        "line.\n\n"
        "Surprisingly, the slope is NEGATIVE — more VRAM is associated with LOWER "
        "premiums. The coefficient is negative two point seven six, and significant. "
        "This contradicts everything we'd expect.\n\n"
        "[Pause]\n\n"
        "We've been tricked by MSRP. MSRP is a confounding variable. High-VRAM cards "
        "like the RTX 3090 — 24 gigabytes — launched at fourteen ninety-nine. The "
        "RTX 3060 with 12 gigabytes launched at just three twenty-nine. A high MSRP "
        "limits the percentage premium you can charge. You simply can't get a three "
        "hundred percent markup on a fifteen hundred dollar card."
    ),

    (
        "Slide 11 — Premium Drivers: Controlling for MSRP",
        "So we run Model 2, controlling for MSRP. The left panel is an added-variable "
        "plot: the X-axis is VRAM after removing MSRP's influence, the Y-axis is "
        "premium after removing MSRP's influence.\n\n"
        "Once we strip out MSRP, the relationship FLIPS positive. More VRAM genuinely "
        "means more premium — after proper control.\n\n"
        "The right panel confirms this: VRAM coefficient is positive two point three "
        "percentage points per gigabyte. MSRP: negative eight points per hundred "
        "dollars. Both significant. MSRP was the hidden confounder — once controlled, "
        "the true relationship emerges."
    ),

    (
        "Slide 12 — Premium Drivers: Adding Tier",
        "Let's go one step further — what if we also account for GPU tier?\n\n"
        "This is Model 3. Look at the premium by tier: XX60 — the entry level — averages "
        "one hundred seven percent premium. That's more than double MSRP. XX70: about "
        "seventy percent. XX80: sixty percent. XX90, the flagship: only fifty-two "
        "percent.\n\n"
        "XX60 has the HIGHEST premium! And when we add Tier into the regression, VRAM "
        "becomes insignificant. Tier is the real driver, not VRAM.\n\n"
        "Which raises the obvious question: why? Why are the cheapest cards getting "
        "the biggest percentage markup?"
    ),

    (
        "Slide 13 — Mining Economics: Cost Per Hashrate",
        "The answer is mining economics — specifically, dollars per megahash per second. "
        "This tells a miner what they pay per unit of mining power. Lower is better.\n\n"
        "Look at the XX60 line at the bottom — it consistently has the lowest cost per "
        "unit of mining power. XX60 cards offer the highest mining return on investment. "
        "So miners bid them up the most in percentage terms.\n\n"
        "This is the key insight: miners didn't target the most powerful cards — they "
        "targeted the most cost-efficient ones. It's not about raw power; it's about "
        "dollars per hash."
    ),

    (
        "Slide 14 — LHR: Background",
        "Now let's talk LHR. In May 2021, at the peak of the mining craze, NVIDIA "
        "announced: 'We're reducing the hash rate of newly manufactured RTX 3080, 3070, "
        "and 3060 Ti cards so they're less desirable to miners.'\n\n"
        "The question: did it actually work?"
    ),

    (
        "Slide 15 — LHR: Real Impact",
        "Here's the comparison — LHR versus non-LHR for the same models. All real "
        "published prices.\n\n"
        "During the mining boom of mid-2021, the gap was big: LHR cards sold twenty "
        "to thirty-seven percent below non-LHR. Miners clearly valued full hash rate.\n\n"
        "But here's the fascinating part: after the crypto crash in mid-2022, the gap "
        "nearly vanished — narrowing to about three percent.\n\n"
        "The LHR price gap is NOT constant — it varies with market conditions. LHR only "
        "matters when mining is profitable. When mining isn't profitable, nobody cares "
        "about hash rate, and prices converge."
    ),

    (
        "Slide 16 — Mining Economics With LHR",
        "From the mining economics angle: dashed lines are LHR cards. They consistently "
        "sit above their non-LHR counterparts — LHR variants have a forty to fifty "
        "percent higher cost per unit of mining power.\n\n"
        "NVIDIA succeeded. They made mining on LHR cards far less rewarding. The "
        "technology worked exactly as intended."
    ),

    # ===== SECTION 3: CONCLUSION (Slides 17-21, ~1.5 min, ~200 words) =====

    (
        "Slide 17 — Answers to Research Questions",
        "So, to our research questions. Do crypto prices drive GPU secondary-market "
        "prices? Yes — and the connection is quite strong, confirmed by both visual "
        "evidence and first-difference correlation tests.\n\n"
        "Which GPU attributes explain the premium? It's mining cost-efficiency — "
        "dollars per megahash — which explains why entry-level cards had the highest "
        "percentage premiums.\n\n"
        "Do different tiers show different volatility? Not really — all four tiers "
        "show nearly identical month-over-month volatility.\n\n"
        "Did LHR work? Yes, and quite successfully — a twenty to thirty-seven percent "
        "price gap during the boom, though the effect disappeared when mining became "
        "unprofitable."
    ),

    (
        "Slide 18 — Limitations",
        "We have limitations: ten GPU models over twenty-four months, eBay as the sole "
        "resale channel, LHR prices rely on range midpoints, and we can't control for "
        "external factors like the pandemic, chip shortage, or scalper activity — all "
        "of which coincided with our study period."
    ),

    (
        "Slide 19 — Future Work",
        "Future directions: the 2022 LHR unlock events provide a natural experiment "
        "for difference-in-differences design. Multi-country comparisons across US, "
        "European, and Asian second-hand markets could test whether findings generalize. "
        "And we could try to disentangle the pandemic and chip shortage effects."
    ),

    (
        "Slide 20 — References",
        "We acknowledge our data sources: CoinGecko, Tom's Hardware, PCMag, 3D Center, "
        "and TechSpot. Full academic references are in the paper. AI assistance via "
        "DeepSeek V4 Pro Max through Claude Code was used for data collection scripts."
    ),

    (
        "Slide 21 — Thank You",
        "Thank you. I'm happy to take any questions."
    ),
]

# --- Write to docx ---
title = doc.add_heading("Presentation Script", level=0)
subtitle = doc.add_paragraph(
    "Analysis of NVIDIA GPU Market Dynamics and Cryptocurrency Correlations\n"
    "2020/09 – 2022/09: A Full-Cycle Study from Mining Boom to Market Crash"
)
subtitle.style = doc.styles["Subtitle"]

doc.add_paragraph("Presenter: YAN Weibo")
doc.add_paragraph("Estimated speaking time: ~10 minutes")
doc.add_paragraph(
    "Note: Slides 8–16 (Results) are the core of the presentation and should receive "
    "the most time. Adjust pacing as needed to stay within 10 minutes."
).runs[0].italic = True

doc.add_paragraph("—" * 40)

for label, text in slides:
    heading = doc.add_heading(label, level=2)
    para = doc.add_paragraph(text)
    # Add a small separator between slides
    doc.add_paragraph("")

# --- Save ---
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print(f"Script saved to: {OUT_PATH}")

# Word count
total_words = sum(len(t.split()) for _, t in slides)
print(f"Total word count: {total_words}")
print(f"Estimated speaking time at 140 wpm: {total_words / 140:.1f} minutes")
